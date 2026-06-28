"""설교 앱 백엔드 - FastAPI + Whisper + Claude + DOCX/PDF 생성"""

import os
import re
import tempfile
from datetime import datetime
from pathlib import Path

from dotenv import load_dotenv
from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel

load_dotenv()

app = FastAPI(title="설교 앱 백엔드", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

SERMONS_DIR = Path("sermons")
SERMONS_DIR.mkdir(exist_ok=True)

# Lazy-loaded Whisper model
_whisper_model = None


def _get_whisper():
    global _whisper_model
    if _whisper_model is None:
        from faster_whisper import WhisperModel
        _whisper_model = WhisperModel("base", device="cpu", compute_type="int8")
    return _whisper_model


# ── Schemas ──────────────────────────────────────────────────────────────────

class OrganizeRequest(BaseModel):
    title: str
    transcription: str


class SaveRequest(BaseModel):
    title: str
    date: str
    transcription: str
    organized_text: str


# ── Endpoints ─────────────────────────────────────────────────────────────────

@app.get("/health")
def health():
    return {"status": "ok"}


@app.post("/transcribe")
async def transcribe_audio(audio: UploadFile = File(...)):
    """음성 파일을 텍스트로 변환 (STT)"""
    suffix = Path(audio.filename or "audio.m4a").suffix or ".m4a"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(await audio.read())
        tmp_path = tmp.name

    try:
        model = _get_whisper()
        segments, _ = model.transcribe(tmp_path, language="ko")
        text = " ".join(s.text.strip() for s in segments)
        return {"transcription": text.strip()}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"STT 오류: {e}")
    finally:
        os.unlink(tmp_path)


@app.post("/organize")
async def organize_sermon(req: OrganizeRequest):
    """Claude AI로 설교 정리"""
    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        raise HTTPException(
            status_code=500,
            detail="ANTHROPIC_API_KEY 환경변수가 설정되지 않았습니다. backend/.env 파일을 확인하세요.",
        )

    import anthropic

    client = anthropic.Anthropic(api_key=api_key)
    prompt = f"""다음은 설교 녹음을 텍스트로 변환한 내용입니다. 체계적인 설교문 형식으로 정리해 주세요.

설교 제목: {req.title}

원본 녹취:
{req.transcription}

아래 형식으로 깔끔하게 정리해 주세요:

【설교 제목】
{req.title}

【서론】
(도입부 내용)

【본문】
(성경 본문 및 핵심 메시지를 단락별로)

【결론 및 적용】
(결론과 성도들이 삶에 적용할 내용)"""

    try:
        msg = client.messages.create(
            model="claude-opus-4-8",
            max_tokens=4096,
            messages=[{"role": "user", "content": prompt}],
        )
        return {"organized_text": msg.content[0].text}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI 정리 오류: {e}")


@app.post("/save")
async def save_sermon(req: SaveRequest):
    """DOCX 및 PDF 파일 생성 후 저장"""
    try:
        date_obj = datetime.fromisoformat(req.date)
    except ValueError:
        date_obj = datetime.now()

    date_str = date_obj.strftime("%Y%m%d")
    safe = re.sub(r'[\\/*?:"<>|]', "", req.title).strip().replace(" ", "_")
    filename = f"{date_str}_{safe}"

    try:
        _build_docx(filename, req, date_obj)
        _build_pdf(filename, req, date_obj)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"파일 생성 오류: {e}")

    return {
        "filename": filename,
        "docx_file": f"{filename}.docx",
        "pdf_file": f"{filename}.pdf",
    }


@app.get("/download/{filename:path}")
async def download(filename: str):
    """파일 다운로드"""
    path = SERMONS_DIR / filename
    if not path.exists():
        raise HTTPException(status_code=404, detail="파일을 찾을 수 없습니다")
    return FileResponse(str(path), filename=filename)


@app.get("/sermons")
async def list_sermons():
    """저장된 설교 목록"""
    seen: dict[str, dict] = {}
    for f in sorted(SERMONS_DIR.iterdir(), reverse=True):
        if f.suffix not in (".docx", ".pdf"):
            continue
        stem = f.stem
        if stem not in seen:
            seen[stem] = {"filename": stem, "docx": None, "pdf": None}
            try:
                seen[stem]["date"] = f"{stem[:4]}-{stem[4:6]}-{stem[6:8]}"
                seen[stem]["title"] = stem[9:].replace("_", " ")
            except Exception:
                seen[stem]["date"] = ""
                seen[stem]["title"] = stem
        seen[stem]["docx" if f.suffix == ".docx" else "pdf"] = f.name

    return {"sermons": list(seen.values())}


# ── File builders ─────────────────────────────────────────────────────────────

def _build_docx(filename: str, req: SaveRequest, date_obj: datetime):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(11)

    h = doc.add_heading(req.title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    dp = doc.add_paragraph(date_obj.strftime("%Y년 %m월 %d일"))
    dp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_heading("정리된 설교문", level=1)
    doc.add_paragraph(req.organized_text)

    doc.add_page_break()
    doc.add_heading("원본 녹취록", level=1)
    doc.add_paragraph(req.transcription)

    doc.save(str(SERMONS_DIR / f"{filename}.docx"))


def _build_pdf(filename: str, req: SaveRequest, date_obj: datetime):
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

    # Korean CID font bundled with ReportLab
    pdfmetrics.registerFont(UnicodeCIDFont("HYSMyeongJoStd-Medium"))
    KO = "HYSMyeongJoStd-Medium"

    title_s = ParagraphStyle("T", fontName=KO, fontSize=18, alignment=TA_CENTER, spaceAfter=4)
    date_s = ParagraphStyle("D", fontName=KO, fontSize=11, alignment=TA_CENTER, spaceAfter=16, textColor="#666666")
    h1_s = ParagraphStyle("H", fontName=KO, fontSize=14, spaceBefore=12, spaceAfter=6, fontWeight="bold")
    body_s = ParagraphStyle("B", fontName=KO, fontSize=11, leading=20, spaceAfter=6)

    def para(text: str, style: ParagraphStyle) -> Paragraph:
        return Paragraph(text.replace("\n", "<br/>"), style)

    doc = SimpleDocTemplate(
        str(SERMONS_DIR / f"{filename}.pdf"),
        pagesize=A4,
        leftMargin=22 * mm,
        rightMargin=22 * mm,
        topMargin=22 * mm,
        bottomMargin=22 * mm,
    )

    doc.build([
        para(req.title, title_s),
        para(date_obj.strftime("%Y년 %m월 %d일"), date_s),
        Spacer(1, 8),
        para("정리된 설교문", h1_s),
        para(req.organized_text, body_s),
        PageBreak(),
        para("원본 녹취록", h1_s),
        para(req.transcription, body_s),
    ])
